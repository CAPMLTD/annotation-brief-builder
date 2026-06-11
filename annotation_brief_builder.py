import streamlit as st
import anthropic
import os
import json
import base64
from pathlib import Path
from datetime import datetime
from io import BytesIO

# ── Page config ────────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="CA Annotation Brief Builder",
    page_icon="📋",
    layout="wide",
    initial_sidebar_state="collapsed",
)

# ── Styling ────────────────────────────────────────────────────────────────────
st.markdown("""
<style>
  @import url('https://fonts.googleapis.com/css2?family=DM+Sans:wght@300;400;500;600;700&family=DM+Mono:wght@400;500&display=swap');

  html, body, [class*="css"] { font-family: 'DM Sans', sans-serif; }
  .main { background: #f8f9fb; }

  .hero {
    background: linear-gradient(135deg, #0f172a 0%, #0d3d2a 60%, #0a9b5c 100%);
    border-radius: 16px;
    padding: 2rem 3rem;
    margin-bottom: 2rem;
    color: white;
  }
  .hero-inner { display: flex; align-items: center; gap: 1.5rem; }
  .hero-logo { width: 72px; height: 72px; border-radius: 10px; background: white; padding: 6px; flex-shrink: 0; }
  .hero-text h1 { font-size: 1.85rem; font-weight: 700; margin: 0 0 0.3rem 0; letter-spacing: -0.5px; color: white; }
  .hero-text p  { font-size: 0.95rem; opacity: 0.8; margin: 0 0 0.4rem 0; }
  .hero-tagline { font-size: 0.75rem; opacity: 0.6; letter-spacing: 1px; text-transform: uppercase; font-weight: 500; color: #a7f3d0; }

  .step-indicator {
    display: flex;
    gap: 0.5rem;
    margin-bottom: 1.5rem;
    flex-wrap: wrap;
  }
  .step-pill {
    padding: 0.3rem 0.8rem;
    border-radius: 999px;
    font-size: 0.75rem;
    font-weight: 600;
    border: 1.5px solid #e2e8f0;
    color: #94a3b8;
    background: white;
  }
  .step-pill.active {
    background: #0a9b5c;
    color: white;
    border-color: #0a9b5c;
  }
  .step-pill.complete {
    background: #dcfce7;
    color: #166534;
    border-color: #0a9b5c;
  }

  .section-header {
    background: #0f172a;
    color: white;
    padding: 0.75rem 1.25rem;
    border-radius: 8px 8px 0 0;
    font-weight: 600;
    font-size: 0.95rem;
    letter-spacing: 0.3px;
    margin-top: 1.5rem;
    border-left: 4px solid #0a9b5c;
  }
  .section-body {
    border: 1px solid #e2e8f0;
    border-top: none;
    border-radius: 0 0 8px 8px;
    padding: 1.5rem;
    background: white;
    margin-bottom: 0.5rem;
  }

  .use-case-card {
    background: white;
    border: 1.5px solid #e2e8f0;
    border-radius: 10px;
    padding: 1rem 1.25rem;
    margin-bottom: 0.75rem;
    border-left: 4px solid #0a9b5c;
  }

  .confirmability-score {
    border-radius: 8px;
    padding: 0.75rem 1rem;
    margin: 0.75rem 0;
    font-size: 0.85rem;
  }
  .score-high { background: #dcfce7; border: 1px solid #0a9b5c; color: #166534; }
  .score-medium { background: #fef9c3; border: 1px solid #eab308; color: #854d0e; }
  .score-low { background: #fee2e2; border: 1px solid #dc2626; color: #991b1b; }

  .stress-test-card {
    background: #f8fafc;
    border: 1.5px solid #e2e8f0;
    border-radius: 8px;
    padding: 1rem;
    margin-bottom: 0.75rem;
  }

  .quality-bar-outer {
    background: #e2e8f0;
    border-radius: 999px;
    height: 10px;
    margin: 0.4rem 0;
  }
  .quality-bar-inner {
    height: 10px;
    border-radius: 999px;
    background: linear-gradient(90deg, #0a9b5c, #0d3d2a);
  }

  .stTextInput > div > div > input,
  .stTextArea > div > div > textarea,
  .stSelectbox > div > div {
    border-radius: 8px !important;
    border: 1.5px solid #e2e8f0 !important;
    font-family: 'DM Sans', sans-serif !important;
  }
  .stTextInput > div > div > input:focus,
  .stTextArea > div > div > textarea:focus {
    border-color: #0a9b5c !important;
    box-shadow: 0 0 0 3px rgba(10,155,92,0.15) !important;
  }
  .stButton > button {
    border-radius: 8px !important;
    font-family: 'DM Sans', sans-serif !important;
    font-weight: 600 !important;
  }

  .generated-brief {
    background: white;
    border: 1px solid #e2e8f0;
    border-radius: 12px;
    padding: 2rem;
    margin-top: 1.5rem;
  }

  .quality-dimension {
    display: flex;
    align-items: center;
    justify-content: space-between;
    padding: 0.5rem 0;
    border-bottom: 1px solid #f1f5f9;
  }
</style>
""", unsafe_allow_html=True)

# ── Logo ───────────────────────────────────────────────────────────────────────
def get_logo_base64():
    logo_path = Path("logo.png")
    if logo_path.exists():
        with open(logo_path, "rb") as f:
            return base64.b64encode(f.read()).decode()
    return None

logo_b64 = get_logo_base64()
logo_html = f'<img src="data:image/png;base64,{logo_b64}" class="hero-logo" alt="CA Logo"/>' if logo_b64 else '<div class="hero-logo" style="display:flex;align-items:center;justify-content:center;font-weight:700;font-size:1.5rem;color:#0a9b5c;">CA</div>'

# ── Helpers ────────────────────────────────────────────────────────────────────
def section_header(title, subtitle=None):
    sub = f'<span style="font-size:0.75rem;opacity:0.7;font-weight:400;margin-left:0.5rem;">{subtitle}</span>' if subtitle else ""
    st.markdown(f'<div class="section-header">{title}{sub}</div>', unsafe_allow_html=True)
    st.markdown('<div class="section-body">', unsafe_allow_html=True)

def section_end():
    st.markdown('</div>', unsafe_allow_html=True)

STEPS = [
    "Project Setup",
    "Data & Privacy",
    "Environment & Glossary",
    "Annotator Onboarding",
    "Use Case Builder",
    "Mechanics & QC",
    "Governance",
    "Generate Brief",
]

def step_indicator(current):
    pills = ""
    for i, s in enumerate(STEPS):
        if i < current:
            cls = "complete"
            label = f"✓ {s}"
        elif i == current:
            cls = "active"
            label = s
        else:
            cls = ""
            label = s
        pills += f'<span class="step-pill {cls}">{label}</span>'
    st.markdown(f'<div class="step-indicator">{pills}</div>', unsafe_allow_html=True)

def get_client():
    api_key = os.environ.get("ANTHROPIC_API_KEY", "")
    if not api_key:
        st.error("ANTHROPIC_API_KEY not set. Please set it before launching.")
        st.stop()
    return anthropic.Anthropic(api_key=api_key)

MODALITY_OPTIONS = [
    "Computer Vision (images / video)",
    "Natural Language Processing (text)",
    "Audio",
    "Multimodal (e.g. image + text)",
    "Other / Custom",
]

# ── Session state init ─────────────────────────────────────────────────────────
defaults = {
    "step": 0,
    "d": {},
    "use_cases": [],
    "current_uc": {},
    "uc_editing": False,
    "stress_tests": {},
    "stress_answers": {},
    "generated_brief": "",
    "quality_scores": {},
}
for k, v in defaults.items():
    if k not in st.session_state:
        st.session_state[k] = v

d = st.session_state.d

# ── Hero ───────────────────────────────────────────────────────────────────────
st.markdown(f"""
<div class="hero">
  <div class="hero-inner">
    {logo_html}
    <div class="hero-text">
      <h1>📋 Annotation Brief Builder</h1>
      <p>Build comprehensive, professional annotation briefs with built-in Confirmability Scoring, Boundary Case Stress Testing, and Brief Quality Assessment — powered by Claude</p>
      <span class="hero-tagline">Excellence | Honesty | Quality &nbsp;·&nbsp; CA Project Management Services Ltd &nbsp;·&nbsp; caprojectmgmt.com</span>
    </div>
  </div>
</div>
""", unsafe_allow_html=True)

step = st.session_state.step
step_indicator(step)

# ══════════════════════════════════════════════════════════════════════════════
# STEP 0 — PROJECT SETUP
# ══════════════════════════════════════════════════════════════════════════════
if step == 0:
    st.subheader("Step 1 — Project Setup")

    section_header("1 · Project Overview")
    col1, col2 = st.columns(2)
    with col1:
        d["project_name"]  = st.text_input("Project name *", value=d.get("project_name",""), placeholder="e.g. Retail SCO Shrink Detection — Phase 2")
        d["project_owner"] = st.text_input("Project owner *", value=d.get("project_owner",""), placeholder="e.g. Head of ML / Annotation Lead")
        d["organisation"]  = st.text_input("Organisation *", value=d.get("organisation",""), placeholder="e.g. Acme AI Ltd")
    with col2:
        d["timeline"]      = st.text_input("Annotation timeline", value=d.get("timeline",""), placeholder="e.g. 3 weeks, starting 01/06/2026")
        d["annotation_tool"] = st.text_input("Annotation tool / platform", value=d.get("annotation_tool",""), placeholder="e.g. Label Studio, CVAT, Prodigy, Scale AI")
        d["brief_version"] = st.text_input("Brief version", value=d.get("brief_version","1.0"), placeholder="e.g. 1.0")
    d["task_description"] = st.text_area("Task description *", value=d.get("task_description",""), height=100,
        placeholder="What is being annotated, why, and what ML task does it support? What is the intended model use case?")
    d["annotation_objectives"] = st.text_area("Annotation objectives *", value=d.get("annotation_objectives",""), height=80,
        placeholder="What does the annotation exercise need to produce? What quality standard must it meet?")
    section_end()

    section_header("2 · Data Modality")
    d["modality"] = st.selectbox("Data modality *", MODALITY_OPTIONS,
        index=MODALITY_OPTIONS.index(d.get("modality", MODALITY_OPTIONS[0])))
    if d["modality"] == "Other / Custom":
        d["modality_custom"] = st.text_input("Describe your data modality", value=d.get("modality_custom",""))

    st.markdown("**The Confirmability Principle** — the governing rule for this brief:")
    if "computer vision" in d["modality"].lower():
        default_confirmability = "Only label what is visually observable and confirmable from the source footage or image alone. Do not infer intent, context, or off-screen events. If you cannot confirm what you are seeing from the data provided, do not label — apply the Confirmability Rule for the relevant use case."
    elif "nlp" in d["modality"].lower() or "language" in d["modality"].lower():
        default_confirmability = "Only label what is explicitly present in the text. Do not infer intent, implied meaning, or context beyond what is written. If the label cannot be confirmed from the text span alone, do not label — apply the Confirmability Rule for the relevant use case."
    elif "audio" in d["modality"].lower():
        default_confirmability = "Only label what is clearly audible and confirmable from the audio segment alone. Do not infer speaker intent or off-recording context. If the label cannot be confirmed from the audio provided, do not label — apply the Confirmability Rule for the relevant use case."
    else:
        default_confirmability = "Only label what is directly observable and confirmable from the source data alone. Do not infer intent or context beyond what is present. If you cannot confirm the label from the data provided, do not label — apply the Confirmability Rule for the relevant use case."

    d["confirmability_principle"] = st.text_area(
        "Confirmability Principle (edit to suit your project) *",
        value=d.get("confirmability_principle", default_confirmability),
        height=120
    )
    section_end()

    st.markdown("---")
    if st.button("Continue to Data & Privacy →", type="primary", use_container_width=True):
        if not d.get("project_name") or not d.get("task_description") or not d.get("annotation_objectives"):
            st.error("Please complete all required fields (*) before continuing.")
        else:
            st.session_state.step = 1
            st.rerun()

# ══════════════════════════════════════════════════════════════════════════════
# STEP 1 — DATA & PRIVACY
# ══════════════════════════════════════════════════════════════════════════════
elif step == 1:
    st.subheader("Step 2 — Data & Privacy")

    section_header("3 · Data Description")
    col1, col2 = st.columns(2)
    with col1:
        d["data_sources"]  = st.text_area("Data sources *", value=d.get("data_sources",""), height=100,
            placeholder="Where does the source data come from? Proprietary capture, third-party datasets, web-scraped, synthetic?")
        d["data_volume"]   = st.text_input("Data volume *", value=d.get("data_volume",""), placeholder="e.g. 10,000 video clips / 50,000 text spans / 2,000 audio segments")
        d["data_format"]   = st.text_input("Data format", value=d.get("data_format",""), placeholder="e.g. MP4 clips at 1080p 25fps / JSON text / WAV 16kHz mono")
    with col2:
        d["data_distribution"] = st.text_area("Sample distribution", value=d.get("data_distribution",""), height=100,
            placeholder="How is the data distributed across classes, environments, conditions? Known imbalances?")
        d["data_quality_issues"] = st.text_area("Known quality issues", value=d.get("data_quality_issues",""), height=100,
            placeholder="Occlusion, low resolution, background noise, transcription errors, missing segments?")
    d["data_date_range"] = st.text_input("Data date range", value=d.get("data_date_range",""), placeholder="e.g. Jan 2025 – Mar 2026")
    section_end()

    section_header("4 · Data Privacy & Handling")
    col1, col2 = st.columns(2)
    with col1:
        d["data_classification"] = st.selectbox("Data classification *",
            ["Select...", "Public", "Internal", "Confidential", "Restricted / Personal Data"],
            index=["Select...", "Public", "Internal", "Confidential", "Restricted / Personal Data"].index(d.get("data_classification","Select...")))
        d["gdpr_basis"] = st.text_input("GDPR lawful basis (if personal data)", value=d.get("gdpr_basis",""),
            placeholder="e.g. Legitimate interests / Consent / Contract")
        d["data_storage"] = st.text_area("Data storage requirements", value=d.get("data_storage",""), height=80,
            placeholder="Where must data be stored? Access controls? Encryption requirements?")
    with col2:
        d["annotator_permitted"] = st.text_area("Annotators MAY *", value=d.get("annotator_permitted",""), height=80,
            placeholder="e.g. Access data via the annotation platform only. Download clips to complete labelling tasks.")
        d["annotator_prohibited"] = st.text_area("Annotators MAY NOT *", value=d.get("annotator_prohibited",""), height=80,
            placeholder="e.g. Share, copy, screenshot, or retain source data. Discuss data contents outside the project.")
        d["retention_policy"] = st.text_input("Data retention policy", value=d.get("retention_policy",""),
            placeholder="e.g. Source data deleted 30 days post-project. Annotations retained for 2 years.")
    section_end()

    st.markdown("---")
    col1, col2 = st.columns([1, 4])
    with col1:
        if st.button("← Back"):
            st.session_state.step = 0
            st.rerun()
    with col2:
        if st.button("Continue to Environment & Glossary →", type="primary", use_container_width=True):
            if not d.get("data_sources") or not d.get("data_volume") or d.get("data_classification","Select...") == "Select...":
                st.error("Please complete all required fields (*) before continuing.")
            else:
                st.session_state.step = 2
                st.rerun()

# ══════════════════════════════════════════════════════════════════════════════
# STEP 2 — ENVIRONMENT & GLOSSARY
# ══════════════════════════════════════════════════════════════════════════════
elif step == 2:
    st.subheader("Step 3 — Environment & Glossary")

    section_header("5 · Environment & Zone Definitions")
    modality = d.get("modality","")
    if "computer vision" in modality.lower():
        st.markdown("*Define the spatial zones and boundaries annotators must reference when labelling.*")
        d["zone_definitions"] = st.text_area("Zone definitions *", value=d.get("zone_definitions",""), height=120,
            placeholder="e.g. Zone 1 — Input area (left of separator bar). Zone 2 — Scan zone (scanner bed surface). Zone 3 — Output/bagging area (right of separator bar and bagging scale).")
        d["zone_diagram_instructions"] = st.text_area("Reference diagram instructions", value=d.get("zone_diagram_instructions",""), height=80,
            placeholder="Describe what reference diagrams should be included. e.g. A labelled overhead diagram of the SCO terminal showing all three zone boundaries.")
        d["camera_context"] = st.text_area("Camera / capture context", value=d.get("camera_context",""), height=80,
            placeholder="Camera angle, field of view limitations, known blind spots, resolution considerations.")
    elif "nlp" in modality.lower() or "language" in modality.lower():
        st.markdown("*Define the text spans, label hierarchy, and boundary rules annotators must apply.*")
        d["span_definitions"] = st.text_area("Span / boundary definitions *", value=d.get("span_definitions",""), height=120,
            placeholder="e.g. Entity spans must include the full noun phrase. Sentence boundaries defined by full stop + capitalisation. Overlapping spans not permitted.")
        d["label_hierarchy"] = st.text_area("Label hierarchy", value=d.get("label_hierarchy",""), height=80,
            placeholder="e.g. Top-level: Sentiment (Positive / Negative / Neutral). Sub-level: Aspect (Product / Service / Delivery).")
        d["zone_definitions"] = d.get("span_definitions","")
    elif "audio" in modality.lower():
        st.markdown("*Define the temporal boundaries and segment rules annotators must apply.*")
        d["temporal_boundaries"] = st.text_area("Temporal boundary rules *", value=d.get("temporal_boundaries",""), height=120,
            placeholder="e.g. Segment start = first audible onset of target sound. Segment end = last audible offset + 0.1s buffer. Silence gaps > 0.5s treated as segment boundaries.")
        d["zone_definitions"] = d.get("temporal_boundaries","")
    else:
        d["zone_definitions"] = st.text_area("Boundary / zone definitions *", value=d.get("zone_definitions",""), height=120,
            placeholder="Define any spatial, temporal, or structural boundaries annotators must apply.")

    d["environment_notes"] = st.text_area("Additional environment notes", value=d.get("environment_notes",""), height=80,
        placeholder="Any other contextual information annotators need to understand the data environment.")
    section_end()

    section_header("6 · Glossary")
    st.markdown("Define controlled vocabulary for terms used in this brief. Add terms that might be ambiguous or interpreted differently by different annotators.")
    if "glossary" not in d:
        d["glossary"] = []

    col1, col2, col3 = st.columns([2, 3, 1])
    with col1:
        new_term = st.text_input("Term", key="new_term", placeholder="e.g. Store product")
    with col2:
        new_def = st.text_input("Definition", key="new_def", placeholder="e.g. Any item with a barcode intended for retail sale, excluding personal belongings and staff equipment.")
    with col3:
        st.markdown("<br>", unsafe_allow_html=True)
        if st.button("Add Term", type="secondary"):
            if new_term and new_def:
                d["glossary"].append({"term": new_term, "definition": new_def})
                st.rerun()

    if d["glossary"]:
        for i, entry in enumerate(d["glossary"]):
            col1, col2, col3 = st.columns([2, 4, 1])
            with col1:
                st.markdown(f"**{entry['term']}**")
            with col2:
                st.markdown(entry['definition'])
            with col3:
                if st.button("✕", key=f"del_term_{i}"):
                    d["glossary"].pop(i)
                    st.rerun()
    else:
        st.info("No glossary terms added yet. Add at least the key domain-specific terms used in your use case rules.")
    section_end()

    st.markdown("---")
    col1, col2 = st.columns([1, 4])
    with col1:
        if st.button("← Back"):
            st.session_state.step = 1
            st.rerun()
    with col2:
        if st.button("Continue to Annotator Onboarding →", type="primary", use_container_width=True):
            if not d.get("zone_definitions"):
                st.error("Please complete the zone/boundary definitions before continuing.")
            else:
                st.session_state.step = 3
                st.rerun()

# ══════════════════════════════════════════════════════════════════════════════
# STEP 3 — ANNOTATOR ONBOARDING
# ══════════════════════════════════════════════════════════════════════════════
elif step == 3:
    st.subheader("Step 4 — Annotator Onboarding")

    section_header("7 · Annotator Onboarding")
    st.markdown("Define the orientation content that annotators receive before starting labelling work.")

    col1, col2 = st.columns(2)
    with col1:
        d["how_to_navigate"] = st.text_area("How to navigate this brief *", value=d.get("how_to_navigate",""), height=100,
            placeholder="e.g. Read Section 1-3 before touching any data. Each use case follows the same format: Summary → Observable Signals → When to Label → Do Not Label → Boundary Cases. Apply the Confirmability Rule before labelling any clip.")
        d["multiple_events"] = st.text_area("How to handle multiple simultaneous events", value=d.get("multiple_events",""), height=80,
            placeholder="e.g. If two anomalies occur in the same clip, label both. Apply separate bounding boxes for each event. Add a comment in the annotation tool noting the multiple event.")
    with col2:
        d["not_covered"] = st.text_area("What to do when a scenario isn't covered *", value=d.get("not_covered",""), height=100,
            placeholder="e.g. Classify as 'Other', add a descriptive comment, and flag for review. Do not attempt to force-fit the clip into an existing use case. Novel scenarios will be reviewed and ruled on by the project lead.")
        d["escalation_overview"] = st.text_area("Escalation pathway overview *", value=d.get("escalation_overview",""), height=80,
            placeholder="e.g. If uncertain: do not label. Flag using the annotation tool's flag function. Do not leave ambiguous clips unlabelled and unflagged. Flagged clips reviewed within 48 hours.")

    d["onboarding_notes"] = st.text_area("Additional onboarding notes", value=d.get("onboarding_notes",""), height=80,
        placeholder="Any other orientation information annotators need before starting work.")
    section_end()

    st.markdown("---")
    col1, col2 = st.columns([1, 4])
    with col1:
        if st.button("← Back"):
            st.session_state.step = 2
            st.rerun()
    with col2:
        if st.button("Continue to Use Case Builder →", type="primary", use_container_width=True):
            if not d.get("how_to_navigate") or not d.get("not_covered") or not d.get("escalation_overview"):
                st.error("Please complete all required fields (*) before continuing.")
            else:
                st.session_state.step = 4
                st.rerun()

# ══════════════════════════════════════════════════════════════════════════════
# STEP 4 — USE CASE BUILDER
# ══════════════════════════════════════════════════════════════════════════════
elif step == 4:
    st.subheader("Step 5 — Use Case Builder")
    st.markdown("Define each annotation class or use case. Add as many as your project requires. Each use case will receive a **Confirmability Score** after you save it.")

    use_cases = st.session_state.use_cases

    # Show existing use cases
    if use_cases:
        st.markdown(f"**{len(use_cases)} use case(s) defined:**")
        for i, uc in enumerate(use_cases):
            score = st.session_state.stress_tests.get(f"score_{i}", None)
            score_html = ""
            if score:
                level = "high" if score >= 7 else "medium" if score >= 4 else "low"
                score_html = f'<span class="confirmability-score score-{level}" style="display:inline-block;padding:0.2rem 0.6rem;border-radius:4px;font-size:0.75rem;font-weight:600;">Confirmability: {score}/10</span>'
            st.markdown(f'<div class="use-case-card"><strong>{uc["name"]}</strong> &nbsp; {score_html}<br><small style="color:#6b7280;">{uc.get("summary","")[:100]}...</small></div>', unsafe_allow_html=True)
            col1, col2 = st.columns([1, 6])
            with col1:
                if st.button("Remove", key=f"remove_uc_{i}"):
                    use_cases.pop(i)
                    st.rerun()

    st.markdown("---")
    st.markdown("### Add a Use Case")

    uc = st.session_state.current_uc

    section_header("Use Case Definition")
    col1, col2 = st.columns(2)
    with col1:
        uc["name"]    = st.text_input("Use case name / class label *", value=uc.get("name",""), placeholder="e.g. Missed Scan / POSITIVE_SENTIMENT / SPEAKER_OVERLAP")
        uc["summary"] = st.text_area("Summary *", value=uc.get("summary",""), height=80,
            placeholder="One paragraph describing what this use case is and when it occurs.")
    with col2:
        uc["observable_signals"] = st.text_area("Observable signals *", value=uc.get("observable_signals",""), height=80,
            placeholder="What can be directly observed in the data that indicates this use case? List the visual/textual/audio signals.")

    col1, col2 = st.columns(2)
    with col1:
        uc["when_to_label"]   = st.text_area("When to label *", value=uc.get("when_to_label",""), height=100,
            placeholder="Define the trigger point. e.g. Label from the frame in which the item crosses the zone boundary. Clip begins 2 seconds before the trigger.")
        uc["label_positive"]  = st.text_area("Label as Positive if *", value=uc.get("label_positive",""), height=100,
            placeholder="List the conditions under which this use case should be labelled. Be specific.")
    with col2:
        uc["do_not_label"]    = st.text_area("Do NOT label if *", value=uc.get("do_not_label",""), height=100,
            placeholder="List the conditions under which this use case must NOT be labelled. This is as important as the positive conditions.")
        uc["boundary_cases"]  = st.text_area("Boundary cases *", value=uc.get("boundary_cases",""), height=100,
            placeholder="Define ambiguous scenarios and the ruling for each. e.g. Item straddles zone boundary — label if majority of item is clearly on the output zone side.")

    col1, col2 = st.columns(2)
    with col1:
        uc["confirmability_rule"] = st.text_area("Confirmability Rule for this use case *", value=uc.get("confirmability_rule",""), height=100,
            placeholder="State the specific confirmability rule for this use case. What must be observable and confirmable before labelling? When in doubt, do not label.")
        uc["evidence_type"] = st.text_area("Evidence type & format", value=uc.get("evidence_type",""), height=80,
            placeholder="e.g. Video clip min 5s. Bounding box on item from point of movement through to resting position.")
    with col2:
        uc["example_positive"]  = st.text_area("Clear positive example description *", value=uc.get("example_positive",""), height=80,
            placeholder="Describe what a clear positive example looks like for this use case.")
        uc["example_negative"]  = st.text_area("Clear negative example description *", value=uc.get("example_negative",""), height=80,
            placeholder="Describe what a clear negative example looks like for this use case.")
        uc["example_boundary"]  = st.text_area("Boundary case example description *", value=uc.get("example_boundary",""), height=80,
            placeholder="Describe a boundary case example and state the explicit ruling.")
    section_end()

    # Save + Score use case
    if st.button("💾 Save Use Case & Generate Confirmability Score", type="primary", use_container_width=True):
        required_uc = ["name","summary","observable_signals","when_to_label","label_positive",
                       "do_not_label","boundary_cases","confirmability_rule",
                       "example_positive","example_negative","example_boundary"]
        missing = [f for f in required_uc if not uc.get(f,"").strip()]
        if missing:
            st.error(f"Please complete all required fields before saving.")
        else:
            with st.spinner("Scoring confirmability rule and generating stress test..."):
                client = get_client()
                score_prompt = f"""You are an expert in annotation quality and the Confirmability Principle.

Review this annotation use case confirmability rule and score it 1-10 on:
- Precision: Is the rule specific enough to resolve edge cases without ambiguity?
- Operationality: Can an annotator apply this rule consistently without expert knowledge?
- Completeness: Does it cover what to do when uncertain?

Use Case: {uc['name']}
Summary: {uc['summary']}
Observable Signals: {uc['observable_signals']}
Confirmability Rule: {uc['confirmability_rule']}
Do Not Label conditions: {uc['do_not_label']}
Boundary Cases: {uc['boundary_cases']}

Respond in JSON only:
{{
  "score": <1-10>,
  "level": "<high|medium|low>",
  "strengths": ["<strength 1>", "<strength 2>"],
  "gaps": ["<gap 1>", "<gap 2>"],
  "recommendation": "<one sentence improvement suggestion>",
  "stress_tests": [
    {{"scenario": "<edge case scenario 1>", "correct_ruling": "<Label / Do Not Label>", "reasoning": "<why>"}},
    {{"scenario": "<edge case scenario 2>", "correct_ruling": "<Label / Do Not Label>", "reasoning": "<why>"}},
    {{"scenario": "<edge case scenario 3>", "correct_ruling": "<Label / Do Not Label>", "reasoning": "<why>"}}
  ]
}}"""
                response = client.messages.create(
                    model="claude-sonnet-4-5",
                    max_tokens=1500,
                    messages=[{"role": "user", "content": score_prompt}]
                )
                raw = response.content[0].text
                clean = raw.replace("```json","").replace("```","").strip()
                try:
                    result = json.loads(clean)
                    idx = len(use_cases)
                    use_cases.append(dict(uc))
                    st.session_state.stress_tests[f"score_{idx}"] = result["score"]
                    st.session_state.stress_tests[f"result_{idx}"] = result
                    st.session_state.stress_tests[f"answers_{idx}"] = ["", "", ""]
                    st.session_state.current_uc = {}
                    st.success(f"Use case '{uc['name']}' saved. Confirmability Score: {result['score']}/10")
                    st.rerun()
                except Exception as e:
                    st.error(f"Scoring error: {e}")

    # Show stress tests for saved use cases
    if use_cases:
        st.markdown("---")
        st.markdown("### 🧪 Boundary Case Stress Tests")
        st.markdown("For each saved use case, rule on the three synthetic boundary cases generated by Claude. Your rulings are checked against your defined rules.")

        for i, uc in enumerate(use_cases):
            result = st.session_state.stress_tests.get(f"result_{i}")
            if not result:
                continue
            score = result["score"]
            level = "high" if score >= 7 else "medium" if score >= 4 else "low"

            with st.expander(f"**{uc['name']}** — Confirmability Score: {score}/10", expanded=(score < 7)):
                st.markdown(f'<div class="confirmability-score score-{level}"><strong>Score: {score}/10</strong><br>{"  ·  ".join(result.get("strengths",[]))}</div>', unsafe_allow_html=True)
                if result.get("gaps"):
                    st.markdown(f"⚠️ **Gaps:** {' · '.join(result.get('gaps',[]))}")
                if result.get("recommendation"):
                    st.markdown(f"💡 **Recommendation:** {result['recommendation']}")

                st.markdown("**Stress Test — rule on each scenario:**")
                answers = st.session_state.stress_tests.get(f"answers_{i}", ["","",""])
                for j, st_case in enumerate(result.get("stress_tests",[])):
                    st.markdown(f'<div class="stress-test-card"><strong>Scenario {j+1}:</strong> {st_case["scenario"]}</div>', unsafe_allow_html=True)
                    answers[j] = st.radio(
                        f"Your ruling — Scenario {j+1}",
                        ["Select ruling...", "Label", "Do Not Label", "Flag for Review"],
                        key=f"stress_{i}_{j}",
                        index=["Select ruling...", "Label", "Do Not Label", "Flag for Review"].index(answers[j]) if answers[j] in ["Select ruling...", "Label", "Do Not Label", "Flag for Review"] else 0
                    )
                st.session_state.stress_tests[f"answers_{i}"] = answers

                if st.button(f"Check Rulings for {uc['name']}", key=f"check_{i}"):
                    all_answered = all(a != "Select ruling..." and a != "" for a in answers)
                    if not all_answered:
                        st.warning("Please rule on all three scenarios before checking.")
                    else:
                        contradictions = []
                        for j, st_case in enumerate(result.get("stress_tests",[])):
                            correct = st_case["correct_ruling"]
                            given = answers[j]
                            if given != correct:
                                contradictions.append(f"Scenario {j+1}: You ruled '{given}' but the correct ruling per your defined rules is '{correct}'. {st_case['reasoning']}")
                        if contradictions:
                            st.error("⚠️ **Contradictions detected:**")
                            for c in contradictions:
                                st.markdown(f"- {c}")
                            st.markdown("Review your use case rules to resolve these contradictions before proceeding.")
                        else:
                            st.success("✅ All rulings consistent with your defined rules.")

    st.markdown("---")
    col1, col2 = st.columns([1, 4])
    with col1:
        if st.button("← Back"):
            st.session_state.step = 3
            st.rerun()
    with col2:
        if st.button("Continue to Mechanics & QC →", type="primary", use_container_width=True):
            if not use_cases:
                st.error("Please define at least one use case before continuing.")
            else:
                st.session_state.step = 5
                st.rerun()

# ══════════════════════════════════════════════════════════════════════════════
# STEP 5 — MECHANICS & QC
# ══════════════════════════════════════════════════════════════════════════════
elif step == 5:
    st.subheader("Step 6 — Mechanics, QC & Escalation")

    section_header("8 · Labelling Mechanics")
    col1, col2 = st.columns(2)
    with col1:
        d["bbox_instructions"] = st.text_area("Bounding box / span / boundary instructions *", value=d.get("bbox_instructions",""), height=100,
            placeholder="e.g. Draw bounding box tightly around the visible portion of the item. Maintain box through to resting position even if temporarily obscured. For NLP: include full noun phrase including determiners.")
        d["clip_timings"] = st.text_area("Clip / segment timing rules", value=d.get("clip_timings",""), height=100,
            placeholder="e.g. Clip begins 2 seconds before trigger event. Clip ends 2 seconds after item is stationary. Minimum total clip duration 8 seconds.")
    with col2:
        d["trigger_definitions"] = st.text_area("Trigger event definitions", value=d.get("trigger_definitions",""), height=100,
            placeholder="e.g. Transaction start = moment PAY button is pressed OR first observed scan event if button press not visible. Payment initiation = moment payment screen first becomes visible.")
        d["file_naming"] = st.text_area("File naming & submission format", value=d.get("file_naming",""), height=100,
            placeholder="e.g. Files named: [ProjectCode]_[UseCase]_[AnnotatorID]_[YYYYMMDD]. Submit completed batches as ZIP via annotation platform by 17:00 each Friday.")
    section_end()

    section_header("9 · Pilot & Calibration Round")
    col1, col2 = st.columns(2)
    with col1:
        d["calibration_set_size"] = st.text_input("Calibration set size *", value=d.get("calibration_set_size",""),
            placeholder="e.g. 50 pre-labelled samples per use case")
        d["calibration_threshold"] = st.text_input("Acceptance threshold *", value=d.get("calibration_threshold",""),
            placeholder="e.g. 85% agreement with reference labels required to proceed to production labelling")
    with col2:
        d["calibration_fail"] = st.text_area("If annotator fails calibration", value=d.get("calibration_fail",""), height=80,
            placeholder="e.g. Review failed examples with project lead. Repeat calibration on a fresh set of 25 samples. Two failures result in removal from the project.")
    d["calibration_notes"] = st.text_area("Calibration notes", value=d.get("calibration_notes",""), height=60,
        placeholder="Any additional notes on the calibration process.")
    section_end()

    section_header("10 · Inter-Annotator Agreement")
    col1, col2 = st.columns(2)
    with col1:
        d["iaa_metric"] = st.selectbox("IAA metric *",
            ["Select...", "Cohen's Kappa", "Fleiss' Kappa", "Percentage Agreement", "Krippendorff's Alpha", "F1 Agreement", "Custom"],
            index=["Select...", "Cohen's Kappa", "Fleiss' Kappa", "Percentage Agreement", "Krippendorff's Alpha", "F1 Agreement", "Custom"].index(d.get("iaa_metric","Select...")))
        d["iaa_target"] = st.text_input("IAA target *", value=d.get("iaa_target",""),
            placeholder="e.g. Cohen's Kappa ≥ 0.80 / Percentage agreement ≥ 90%")
        d["iaa_cadence"] = st.text_input("Measurement cadence *", value=d.get("iaa_cadence",""),
            placeholder="e.g. Measured weekly on a random 10% sample of completed annotations")
    with col2:
        d["iaa_disagreement"] = st.text_area("Disagreement resolution process *", value=d.get("iaa_disagreement",""), height=80,
            placeholder="e.g. Disagreements reviewed by project lead within 48 hours. Ruling communicated to all annotators. Affected labels updated.")
        d["iaa_below_threshold"] = st.text_area("Below threshold escalation *", value=d.get("iaa_below_threshold",""), height=80,
            placeholder="e.g. IAA below threshold triggers mandatory team review session. Production labelling paused until threshold is restored.")
    section_end()

    section_header("11 · Escalation Protocol")
    col1, col2 = st.columns(2)
    with col1:
        d["flag_process"] = st.text_area("Flagging process *", value=d.get("flag_process",""), height=80,
            placeholder="e.g. Use the flag function in the annotation tool. Add a brief comment describing the uncertainty. Do not leave ambiguous samples unlabelled and unflagged.")
        d["review_workflow"] = st.text_area("Review workflow *", value=d.get("review_workflow",""), height=80,
            placeholder="e.g. Flagged samples reviewed by project lead within 48 hours. Ruling documented and communicated to all annotators.")
    with col2:
        d["feedback_loop"] = st.text_area("Feedback loop to annotation team", value=d.get("feedback_loop",""), height=80,
            placeholder="e.g. Weekly ruling digest shared with all annotators covering the week's flagged cases and their rulings.")
        d["novel_scenario"] = st.text_area("Novel scenario handling *", value=d.get("novel_scenario",""), height=80,
            placeholder="e.g. Classify as 'Other', add descriptive comment, flag for review. Brief will be updated with new use case definition if novel scenario recurs.")
    section_end()

    section_header("12 · Quality Control Framework")
    col1, col2 = st.columns(2)
    with col1:
        d["review_sampling"] = st.text_input("Review sampling rate *", value=d.get("review_sampling",""),
            placeholder="e.g. 10% random sample of all completed annotations reviewed by project lead")
        d["rejection_criteria"] = st.text_area("Rejection criteria *", value=d.get("rejection_criteria",""), height=80,
            placeholder="e.g. Batch rejected if error rate exceeds 5%. Annotator placed on review if personal error rate exceeds 8%.")
    with col2:
        d["quality_gates"] = st.text_area("Quality gate thresholds *", value=d.get("quality_gates",""), height=80,
            placeholder="e.g. Phase 1 complete when: IAA ≥ 0.80, error rate < 5%, all use cases have minimum 500 confirmed positive labels.")
        d["remediation"] = st.text_area("Remediation for failed batches", value=d.get("remediation",""), height=80,
            placeholder="e.g. Failed batch returned to annotator with error examples highlighted. Re-labelled batch undergoes 20% review rather than standard 10%.")
    section_end()

    st.markdown("---")
    col1, col2 = st.columns([1, 4])
    with col1:
        if st.button("← Back"):
            st.session_state.step = 4
            st.rerun()
    with col2:
        if st.button("Continue to Governance →", type="primary", use_container_width=True):
            required = ["bbox_instructions","calibration_set_size","calibration_threshold",
                        "iaa_target","iaa_cadence","iaa_disagreement","iaa_below_threshold",
                        "flag_process","review_workflow","novel_scenario",
                        "review_sampling","rejection_criteria","quality_gates"]
            missing = [f for f in required if not d.get(f,"").strip()]
            if missing or d.get("iaa_metric","Select...") == "Select...":
                st.error("Please complete all required fields (*) before continuing.")
            else:
                st.session_state.step = 6
                st.rerun()

# ══════════════════════════════════════════════════════════════════════════════
# STEP 6 — GOVERNANCE
# ══════════════════════════════════════════════════════════════════════════════
elif step == 6:
    st.subheader("Step 7 — Governance")

    section_header("13 · Version Control & Amendment Protocol")
    col1, col2 = st.columns(2)
    with col1:
        d["version_control"] = st.text_area("Version control approach *", value=d.get("version_control",""), height=80,
            placeholder="e.g. Brief versioned using Major.Minor. Major version = new use case or fundamental rule change. Minor version = clarification or boundary case addition.")
        d["amendment_triggers"] = st.text_area("What triggers an amendment *", value=d.get("amendment_triggers",""), height=80,
            placeholder="e.g. Novel scenario identified in production. IAA analysis reveals systematic misinterpretation of a rule. Stakeholder review identifies gap.")
    with col2:
        d["amendment_communication"] = st.text_area("How amendments are communicated mid-exercise *", value=d.get("amendment_communication",""), height=80,
            placeholder="e.g. All annotators notified via annotation platform message within 24 hours. Version changelog shared. Briefing call held for Major version changes.")
        d["relabelling_policy"] = st.text_area("Re-labelling policy on rule change *", value=d.get("relabelling_policy",""), height=80,
            placeholder="e.g. Minor version changes: previously labelled data reviewed on a 20% sample only. Major version changes: full re-labelling of affected use case required.")
    section_end()

    section_header("14 · Governance & Sign-Off")
    col1, col2 = st.columns(2)
    with col1:
        d["brief_owner"]    = st.text_input("Brief owner *", value=d.get("brief_owner",""), placeholder="e.g. ML Lead / Annotation Manager")
        d["signoff_authority"] = st.text_input("Sign-off authority *", value=d.get("signoff_authority",""), placeholder="e.g. Head of ML / Project Sponsor")
        d["review_cadence"] = st.text_input("Review cadence", value=d.get("review_cadence",""), placeholder="e.g. Reviewed weekly during active annotation. Formal review at project close.")
    with col2:
        d["document_status"] = st.selectbox("Document status *",
            ["Draft", "Under Review", "Approved", "Active", "Superseded"],
            index=["Draft", "Under Review", "Approved", "Active", "Superseded"].index(d.get("document_status","Draft")))
        d["signoff_date"]   = st.text_input("Sign-off date", value=d.get("signoff_date",""), placeholder=f"e.g. {datetime.now().strftime('%d/%m/%Y')}")
        d["next_review"]    = st.text_input("Next review date", value=d.get("next_review",""), placeholder="e.g. End of annotation phase / 4 weeks from sign-off")
    d["additional_notes"] = st.text_area("Additional notes", value=d.get("additional_notes",""), height=60,
        placeholder="Any final notes, dependencies, or acknowledgements.")
    section_end()

    st.markdown("---")
    col1, col2 = st.columns([1, 4])
    with col1:
        if st.button("← Back"):
            st.session_state.step = 5
            st.rerun()
    with col2:
        if st.button("✨ Generate Annotation Brief", type="primary", use_container_width=True):
            required = ["version_control","amendment_triggers","amendment_communication",
                        "relabelling_policy","brief_owner","signoff_authority"]
            missing = [f for f in required if not d.get(f,"").strip()]
            if missing:
                st.error("Please complete all required fields (*) before generating.")
            else:
                st.session_state.step = 7
                st.rerun()

# ══════════════════════════════════════════════════════════════════════════════
# STEP 7 — GENERATE
# ══════════════════════════════════════════════════════════════════════════════
elif step == 7:
    st.subheader("Step 8 — Generated Annotation Brief")

    use_cases = st.session_state.use_cases

    col1, col2 = st.columns([1, 3])
    with col1:
        if st.button("← Edit"):
            st.session_state.step = 6
            st.rerun()
    with col2:
        if st.button("🔄 Regenerate"):
            st.session_state.generated_brief = ""
            st.rerun()

    if not st.session_state.generated_brief:
        with st.spinner("Generating your Annotation Brief and Quality Assessment..."):
            client = get_client()

            # Build use case summary for prompt
            uc_summary = ""
            for i, uc in enumerate(use_cases):
                score_result = st.session_state.stress_tests.get(f"result_{i}", {})
                uc_summary += f"""
USE CASE {i+1}: {uc['name']}
Summary: {uc['summary']}
Observable Signals: {uc['observable_signals']}
When to Label: {uc['when_to_label']}
Label as Positive If: {uc['label_positive']}
Do NOT Label If: {uc['do_not_label']}
Boundary Cases: {uc['boundary_cases']}
Confirmability Rule: {uc['confirmability_rule']}
Evidence Type: {uc.get('evidence_type','')}
Confirmability Score: {score_result.get('score','N/A')}/10
Confirmability Gaps: {', '.join(score_result.get('gaps',[])) if score_result.get('gaps') else 'None identified'}
Example - Positive: {uc['example_positive']}
Example - Negative: {uc['example_negative']}
Example - Boundary: {uc['example_boundary']}
"""

            glossary_text = "\n".join([f"- {g['term']}: {g['definition']}" for g in d.get("glossary",[])])

            prompt = f"""You are an expert annotation project lead and ML governance specialist. Generate a comprehensive, professional Annotation Brief based on the following project data.

PROJECT DATA:
Project: {d.get('project_name','')}
Organisation: {d.get('organisation','')}
Owner: {d.get('project_owner','')}
Version: {d.get('brief_version','1.0')}
Modality: {d.get('modality','')}
Task: {d.get('task_description','')}
Objectives: {d.get('annotation_objectives','')}
Timeline: {d.get('timeline','')}
Annotation Tool: {d.get('annotation_tool','')}

CONFIRMABILITY PRINCIPLE:
{d.get('confirmability_principle','')}

DATA:
Sources: {d.get('data_sources','')}
Volume: {d.get('data_volume','')}
Format: {d.get('data_format','')}
Distribution: {d.get('data_distribution','')}
Quality Issues: {d.get('data_quality_issues','')}
Date Range: {d.get('data_date_range','')}

DATA PRIVACY:
Classification: {d.get('data_classification','')}
GDPR Basis: {d.get('gdpr_basis','')}
Storage: {d.get('data_storage','')}
Permitted: {d.get('annotator_permitted','')}
Prohibited: {d.get('annotator_prohibited','')}
Retention: {d.get('retention_policy','')}

ENVIRONMENT:
Zone/Boundary Definitions: {d.get('zone_definitions','')}
Environment Notes: {d.get('environment_notes','')}

GLOSSARY:
{glossary_text if glossary_text else 'No glossary terms defined.'}

ANNOTATOR ONBOARDING:
Navigation: {d.get('how_to_navigate','')}
Multiple Events: {d.get('multiple_events','')}
Not Covered: {d.get('not_covered','')}
Escalation Overview: {d.get('escalation_overview','')}

USE CASES:
{uc_summary}

LABELLING MECHANICS:
Bounding Box/Span Instructions: {d.get('bbox_instructions','')}
Clip/Segment Timings: {d.get('clip_timings','')}
Trigger Definitions: {d.get('trigger_definitions','')}
File Naming: {d.get('file_naming','')}

CALIBRATION:
Set Size: {d.get('calibration_set_size','')}
Threshold: {d.get('calibration_threshold','')}
If Fail: {d.get('calibration_fail','')}

IAA:
Metric: {d.get('iaa_metric','')}
Target: {d.get('iaa_target','')}
Cadence: {d.get('iaa_cadence','')}
Disagreement Resolution: {d.get('iaa_disagreement','')}
Below Threshold: {d.get('iaa_below_threshold','')}

ESCALATION:
Flag Process: {d.get('flag_process','')}
Review Workflow: {d.get('review_workflow','')}
Feedback Loop: {d.get('feedback_loop','')}
Novel Scenarios: {d.get('novel_scenario','')}

QC:
Review Sampling: {d.get('review_sampling','')}
Rejection Criteria: {d.get('rejection_criteria','')}
Quality Gates: {d.get('quality_gates','')}
Remediation: {d.get('remediation','')}

GOVERNANCE:
Version Control: {d.get('version_control','')}
Amendment Triggers: {d.get('amendment_triggers','')}
Amendment Communication: {d.get('amendment_communication','')}
Re-labelling Policy: {d.get('relabelling_policy','')}
Brief Owner: {d.get('brief_owner','')}
Sign-off Authority: {d.get('signoff_authority','')}
Document Status: {d.get('document_status','')}
Sign-off Date: {d.get('signoff_date','')}

OUTPUT REQUIREMENTS:
1. Generate a comprehensive, professional Annotation Brief in clean Markdown
2. Structure it with all sections clearly headed
3. For each use case, include the full per-use-case schema in a consistent format
4. The Confirmability Principle must be prominently stated as the governing rule
5. Do NOT invent or fabricate data — only use what is provided
6. Flag any fields left blank as requiring completion before brief is finalised
7. End with a BRIEF QUALITY ASSESSMENT section containing:
   - Completeness Score (0-100%)
   - Confirmability Rigour Rating (average across use cases, with per-use-case breakdown)
   - Ambiguity Risk Rating (RAG: Green/Amber/Red) with specific flags
   - Example Coverage (whether all three example types are specified per use case)
   - Glossary Coverage (whether domain terms used in rules appear in the glossary)
   - IAA Readiness (whether IAA methodology is sufficiently defined)
   - Overall Brief Quality Score (0-100, weighted composite)
   - Top 3 Priority Actions before brief is distributed to annotators

Generate the full brief now:"""

            response = client.messages.create(
                model="claude-sonnet-4-5",
                max_tokens=6000,
                messages=[{"role": "user", "content": prompt}]
            )
            st.session_state.generated_brief = response.content[0].text

    # Display generated brief
    with st.container():
        st.markdown('<div class="generated-brief">', unsafe_allow_html=True)
        st.markdown(st.session_state.generated_brief)
        st.markdown('</div>', unsafe_allow_html=True)

    # Export
    st.markdown("---")
    st.subheader("Export")
    col1, col2, col3 = st.columns(3)

    with col1:
        md_content = f"# Annotation Brief — {d.get('project_name','Untitled')}\n\n"
        md_content += f"**Generated:** {datetime.now().strftime('%Y-%m-%d %H:%M')}  \n"
        md_content += f"**Version:** {d.get('brief_version','1.0')}  \n"
        md_content += f"**Status:** {d.get('document_status','Draft')}\n\n---\n\n"
        md_content += st.session_state.generated_brief
        st.download_button(
            "⬇️ Download Markdown",
            data=md_content.encode("utf-8"),
            file_name=f"annotation_brief_{d.get('project_name','brief').replace(' ','_')}.md",
            mime="text/markdown",
            use_container_width=True,
        )

    with col2:
        try:
            import markdown as md_lib
            html_body = md_lib.markdown(st.session_state.generated_brief, extensions=["tables","fenced_code"])
        except:
            html_body = f"<pre>{st.session_state.generated_brief}</pre>"

        html_content = f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<title>Annotation Brief — {d.get('project_name','')}</title>
<style>
  body {{ font-family: 'Segoe UI', Arial, sans-serif; max-width: 900px; margin: 40px auto; padding: 0 2rem; color: #1a202c; line-height: 1.7; }}
  h1 {{ color: #0f172a; border-bottom: 3px solid #0a9b5c; padding-bottom: 0.5rem; }}
  h2 {{ color: #0d3d2a; margin-top: 2rem; }}
  h3 {{ color: #374151; }}
  table {{ border-collapse: collapse; width: 100%; margin: 1rem 0; }}
  th, td {{ border: 1px solid #e2e8f0; padding: 0.6rem 1rem; text-align: left; }}
  th {{ background: #0f172a; color: white; }}
  tr:nth-child(even) {{ background: #f8f9fb; }}
  code {{ background: #f1f5f9; padding: 0.2rem 0.4rem; border-radius: 4px; }}
  .meta {{ background: #f0fdf4; border: 1px solid #0a9b5c; border-radius: 8px; padding: 1rem; margin-bottom: 2rem; font-size: 0.9rem; }}
  .footer {{ margin-top: 3rem; padding-top: 1rem; border-top: 1px solid #e2e8f0; font-size: 0.8rem; color: #9ca3af; }}
</style>
</head>
<body>
<div class="meta">
  <strong>Project:</strong> {d.get('project_name','')} &nbsp;|&nbsp;
  <strong>Version:</strong> {d.get('brief_version','1.0')} &nbsp;|&nbsp;
  <strong>Status:</strong> {d.get('document_status','Draft')} &nbsp;|&nbsp;
  <strong>Generated:</strong> {datetime.now().strftime('%Y-%m-%d %H:%M')}
</div>
{html_body}
<div class="footer">Generated by CA Project Management Services Ltd — Annotation Brief Builder — <a href="https://www.caprojectmgmt.com" style="color:#0a9b5c;">caprojectmgmt.com</a></div>
</body>
</html>"""
        st.download_button(
            "⬇️ Download HTML",
            data=html_content.encode("utf-8"),
            file_name=f"annotation_brief_{d.get('project_name','brief').replace(' ','_')}.html",
            mime="text/html",
            use_container_width=True,
        )

    with col3:
        try:
            from docx import Document as DocxDocument
            from docx.shared import Pt, RGBColor
            doc = DocxDocument()
            style = doc.styles["Normal"]
            style.font.name = "Arial"
            style.font.size = Pt(11)
            title_para = doc.add_heading(f"Annotation Brief — {d.get('project_name','')}", 0)
            title_para.runs[0].font.color.rgb = RGBColor(0x0f, 0x17, 0x2a)
            meta = doc.add_paragraph()
            meta.add_run(f"Version: {d.get('brief_version','1.0')}   |   Status: {d.get('document_status','Draft')}   |   Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}").italic = True
            doc.add_paragraph()
            for line in st.session_state.generated_brief.split("\n"):
                line = line.strip()
                if not line:
                    doc.add_paragraph()
                elif line.startswith("### "):
                    doc.add_heading(line[4:], level=3)
                elif line.startswith("## "):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith("# "):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith("- ") or line.startswith("* "):
                    doc.add_paragraph(line[2:], style="List Bullet")
                else:
                    p = doc.add_paragraph()
                    parts = line.split("**")
                    for i, part in enumerate(parts):
                        run = p.add_run(part)
                        if i % 2 == 1:
                            run.bold = True
            doc.add_paragraph()
            footer_para = doc.add_paragraph("Generated by CA Project Management Services Ltd — Annotation Brief Builder — caprojectmgmt.com")
            footer_para.runs[0].font.color.rgb = RGBColor(0x9c, 0xa3, 0xaf)
            footer_para.runs[0].font.size = Pt(9)
            buf = BytesIO()
            doc.save(buf)
            buf.seek(0)
            st.download_button(
                "⬇️ Download Word (.docx)",
                data=buf.getvalue(),
                file_name=f"annotation_brief_{d.get('project_name','brief').replace(' ','_')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )
        except ImportError:
            st.info("Install python-docx for Word export.")

    st.markdown("---")
    st.markdown(
        "<div style='text-align:center;color:#9ca3af;font-size:0.8rem;'>CA Project Management Services Ltd &nbsp;·&nbsp; Annotation Brief Builder &nbsp;·&nbsp; <a href='https://www.caprojectmgmt.com' target='_blank' style='color:#0a9b5c;text-decoration:none;'>caprojectmgmt.com</a> &nbsp;·&nbsp; Powered by Claude</div>",
        unsafe_allow_html=True,
    )
